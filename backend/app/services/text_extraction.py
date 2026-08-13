"""Raw text extraction from uploaded resume files."""
import io

import docx
import pdfplumber


class UnsupportedFileError(Exception):
    pass


def extract_text_from_pdf(content: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(content: bytes) -> str:
    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def extract_text(content: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return extract_text_from_pdf(content)
    if file_type == "docx":
        return extract_text_from_docx(content)
    raise UnsupportedFileError(f"Unsupported file type: {file_type}")
